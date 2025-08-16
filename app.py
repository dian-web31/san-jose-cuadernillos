import streamlit as st
from docx import Document
from docxcompose.composer import Composer
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tempfile
import zipfile
import os

# Define número de documentos por Sesion y Grado
materias_por_sesion = {
    "PRIMARIA 1-5": [4, 3, 4],
    "SECUNDARIA 6 Y 7": [5, 4, 4],
    "SECUNDARIA 8 Y 9": [5, 5, 4],
    "MEDIA 10 Y 11": [4, 6, 5]
}

# Función para crear documento combinando varios archivos en una sección
def crear_documento_con_docs(base_path, lista_paths, seccion_num, output_path):
    base_doc = Document(base_path)
    composer = Composer(base_doc)

    for i, doc_path in enumerate(lista_paths):
        doc_to_append = Document(doc_path)
        composer.append(doc_to_append)

    # Añadir texto o encabezado con la sección si quieres:
    composer.doc.add_paragraph(f"Sección {seccion_num}")

    composer.save(output_path)

# --- Streamlit Interface ---

st.title("Generador automático de documentos por secciones")

# Selección de grado
grado = st.selectbox("Selecciona el grado:", list(materias_por_sesion.keys()))

# Subir varios archivos .docx
archivos_subidos = st.file_uploader(
    "Carga tus archivos .docx para combinar (en orden deseado):",
    type=["docx"],
    accept_multiple_files=True
)

# Ruta a plantilla base (debes ajustar esta ruta o subir el archivo)
ruta_plantilla = "docs/FORMATO_COGNITIVAS.docx"

if st.button("Generar documentos"):
    if not archivos_subidos:
        st.error("Por favor, sube al menos un archivo.")
    else:
        # Guardar archivos subidos temporalmente en disco
        temp_dir = tempfile.TemporaryDirectory()
        rutas_docs = []
        for i, arco in enumerate(archivos_subidos):
            ruta_temp = os.path.join(temp_dir.name, f"doc_{i+1}.docx")
            with open(ruta_temp, "wb") as f:
                f.write(arco.getbuffer())
            rutas_docs.append(ruta_temp)

        seccion = 1
        archivos_generados = []
        documentos_acumulados = []

        materias_sesion = materias_por_sesion[grado]  # lista de cuantos docs por sesión, ej: [4,3,4]

        for i, ruta_doc in enumerate(rutas_docs, 1):
            documentos_acumulados.append(ruta_doc)

            # Define el máximo de documentos para la sesión actual
            if seccion <= len(materias_sesion):
                max_docs = materias_sesion[seccion - 1]
            else:
                max_docs = materias_sesion[-1]

            print(f"Sección: {seccion}, max_docs esta sesion: {max_docs}, documentos acumulados: {len(documentos_acumulados)}")

            # Condición para crear documento cuando se alcance el máximo permitido o sea el último documento
            if len(documentos_acumulados) == max_docs or i == len(rutas_docs):
                output_path = os.path.join(temp_dir.name, f"Documento_Seccion_{seccion}.docx")
                crear_documento_con_docs(ruta_plantilla, documentos_acumulados, seccion, output_path)
                archivos_generados.append(output_path)
                documentos_acumulados = []
                seccion += 1

        # Empaquetar los archivos generados en un ZIP para descarga
        zip_path = os.path.join(temp_dir.name, "documentos_seccionados.zip")
        with zipfile.ZipFile(zip_path, "w") as zipf:
            for file_path in archivos_generados:
                zipf.write(file_path, os.path.basename(file_path))

        with open(zip_path, "rb") as fzip:
            st.download_button(
                label="Descargar documentos generados (ZIP)",
                data=fzip,
                file_name="documentos_seccionados.zip",
                mime="application/zip"
            )

        temp_dir.cleanup()
